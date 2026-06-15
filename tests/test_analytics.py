from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
import pandas as pd

from src.output.analytics import (
    analytics_export_path,
    build_analytics_records,
    date_range_for_period,
    export_analytics_excel,
    filter_records_by_date,
    keyword_counts,
    summary_metrics,
)
from src.output.word_writer import add_digest_block


def _write_json(path: Path, payload: object) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")


def _save_master(path: Path) -> None:
    document = Document()
    add_digest_block(
        document,
        "2026-06-02",
        [
            {
                "title_cn": "日本发布AI政策",
                "summary_cn": "摘要",
                "url": "https://example.com/a?utm_source=mail",
            },
            {
                "title_cn": "半导体投资扩大",
                "summary_cn": "摘要",
                "url": "https://example.com/b",
            },
        ],
    )
    document.paragraphs[1].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.save(path)


def test_build_analytics_records_merges_word_highlight_and_digest_metadata(
    tmp_path: Path, monkeypatch
) -> None:
    import src.output.analytics as analytics
    import src.output.cumulative_exporter as cumulative_exporter

    monkeypatch.setattr(analytics, "DATA_ROOT", tmp_path / "data")
    monkeypatch.setattr(cumulative_exporter, "DATA_ROOT", tmp_path / "data")
    master_path = tmp_path / "master.docx"
    _save_master(master_path)
    _write_json(
        tmp_path / "data" / "output" / "cumulative_digest.json",
        {
            "date_blocks": [
                {
                    "date": "2026-06-02",
                    "items": [
                        {
                            "title_cn": "日本发布AI政策",
                            "url": "https://example.com/a",
                            "source": "Nikkei",
                            "type": "政策",
                            "soft_hard": "软科学",
                            "keywords": ["AI", "政策"],
                        },
                        {
                            "title_cn": "半导体投资扩大",
                            "url": "https://example.com/b",
                            "source": "NEDO Japan",
                            "type": "产业",
                            "soft_hard": "硬科学",
                            "keywords": ["半导体"],
                        },
                        {
                            "title_cn": "已从总Word删除的新闻",
                            "url": "https://example.com/deleted",
                            "source": "MEXT Japan",
                            "type": "技术",
                            "soft_hard": "硬科学",
                            "keywords": ["AI for Science"],
                        },
                    ],
                }
            ]
        },
    )

    records = build_analytics_records(master_path)

    assert len(records) == 2
    assert "已从总Word删除的新闻" not in {record.title for record in records}
    assert records[0].accepted is True
    assert records[0].type == "政策"
    assert records[0].source == "Nikkei"
    assert records[0].keywords == ["AI", "政策"]
    assert summary_metrics(records) == {
        "total": 2,
        "accepted": 1,
        "acceptance_rate": 50.0,
        "unaccepted": 1,
    }


def test_date_range_for_week_and_month() -> None:
    assert date_range_for_period("周度", "2026-06-03") == (
        "2026-06-01",
        "2026-06-07",
    )


def test_daily_final_json_enriches_records_when_cumulative_is_partial(
    tmp_path: Path, monkeypatch
) -> None:
    import src.output.analytics as analytics
    import src.output.cumulative_exporter as cumulative_exporter

    monkeypatch.setattr(analytics, "DATA_ROOT", tmp_path / "data")
    monkeypatch.setattr(cumulative_exporter, "DATA_ROOT", tmp_path / "data")
    master_path = tmp_path / "master.docx"
    _save_master(master_path)
    _write_json(
        tmp_path / "data" / "output" / "cumulative_digest.json",
        {"date_blocks": [{"date": "2026-06-02", "items": []}]},
    )
    _write_json(
        tmp_path / "data" / "output" / "final_digest_20260602.json",
        [
            {
                "title_cn": "日本发布AI政策",
                "url": "https://example.com/a",
                "source": "Nikkei",
                "type": "政策",
                "soft_hard": "软科学",
                "keywords": ["AI"],
            }
        ],
    )

    records = build_analytics_records(master_path)

    first = next(record for record in records if record.title == "日本发布AI政策")
    assert first.source == "Nikkei"
    assert first.type == "政策"


def test_source_falls_back_to_url_domain_when_json_metadata_is_missing(
    tmp_path: Path, monkeypatch
) -> None:
    import src.output.analytics as analytics
    import src.output.cumulative_exporter as cumulative_exporter

    monkeypatch.setattr(analytics, "DATA_ROOT", tmp_path / "data")
    monkeypatch.setattr(cumulative_exporter, "DATA_ROOT", tmp_path / "data")
    document = Document()
    add_digest_block(
        document,
        "2026-06-02",
        [
            {
                "title_cn": "日经来源新闻",
                "summary_cn": "摘要",
                "url": "https://www.nikkei.com/article/mock/",
            },
            {
                "title_cn": "科学日本来源新闻",
                "summary_cn": "摘要",
                "url": "https://sj.jst.go.jp/news/202606/mock.html",
            },
        ],
    )
    master_path = tmp_path / "master.docx"
    document.save(master_path)

    records = build_analytics_records(master_path)

    assert {record.title: record.source for record in records} == {
        "日经来源新闻": "Nikkei",
        "科学日本来源新闻": "Science Japan",
    }


def test_candidate_excel_metadata_fills_word_only_records(
    tmp_path: Path, monkeypatch
) -> None:
    import src.output.analytics as analytics
    import src.output.cumulative_exporter as cumulative_exporter

    monkeypatch.setattr(analytics, "DATA_ROOT", tmp_path / "data")
    monkeypatch.setattr(cumulative_exporter, "DATA_ROOT", tmp_path / "data")
    document = Document()
    add_digest_block(
        document,
        "2026-06-03",
        [
            {
                "title_cn": "日本临床试验证实脑梗塞新药疗效",
                "summary_cn": "摘要",
                "url": "https://www.nikkei.com/article/mock-clinical/",
            }
        ],
    )
    master_path = tmp_path / "master.docx"
    document.save(master_path)
    output_dir = tmp_path / "data" / "output"
    output_dir.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(
        [
            {
                "title_original": "日本临床试验证实脑梗塞新药疗效",
                "url": "https://www.nikkei.com/article/mock-clinical/",
                "source": "Nikkei",
                "suggested_type": "技术",
                "suggested_soft_hard": "硬科学",
                "matched_keywords": "生物医药: 临床试验; 生物医药: 新药",
            }
        ]
    ).to_excel(output_dir / "candidates_20260603_combined.xlsx", index=False)

    records = build_analytics_records(master_path)

    assert records[0].data_source == "word+candidate_metadata"
    assert records[0].type == "技术"
    assert records[0].soft_hard == "硬科学"
    assert records[0].keywords == ["临床试验", "新药"]


def test_title_rules_fill_word_only_records_without_candidate_metadata(
    tmp_path: Path, monkeypatch
) -> None:
    import src.output.analytics as analytics
    import src.output.cumulative_exporter as cumulative_exporter

    monkeypatch.setattr(analytics, "DATA_ROOT", tmp_path / "data")
    monkeypatch.setattr(cumulative_exporter, "DATA_ROOT", tmp_path / "data")
    document = Document()
    add_digest_block(
        document,
        "2026-06-03",
        [
            {
                "title_cn": "富岳模拟揭示中微子变身影响超新星爆发",
                "summary_cn": "摘要",
                "url": "https://www.nikkei.com/article/mock-neutrino/",
            }
        ],
    )
    master_path = tmp_path / "master.docx"
    document.save(master_path)

    records = build_analytics_records(master_path)

    assert records[0].data_source == "word+rules"
    assert records[0].source == "Nikkei"
    assert records[0].type == "技术"
    assert records[0].soft_hard == "硬科学"
    assert "中微子" in records[0].keywords


def test_build_analytics_records_uses_digest_json_only_without_master(
    tmp_path: Path, monkeypatch
) -> None:
    import src.output.analytics as analytics
    import src.output.cumulative_exporter as cumulative_exporter

    monkeypatch.setattr(analytics, "DATA_ROOT", tmp_path / "data")
    monkeypatch.setattr(cumulative_exporter, "DATA_ROOT", tmp_path / "data")
    _write_json(
        tmp_path / "data" / "output" / "cumulative_digest.json",
        {
            "date_blocks": [
                {
                    "date": "2026-06-02",
                    "items": [
                        {
                            "title_cn": "仅存在于JSON的新闻",
                            "url": "https://example.com/json-only",
                            "source": "Nikkei",
                            "type": "政策",
                            "soft_hard": "软科学",
                            "keywords": ["政策"],
                        }
                    ],
                }
            ]
        },
    )

    records = build_analytics_records(None)

    assert [record.title for record in records] == ["仅存在于JSON的新闻"]
    assert date_range_for_period("月度", "2026-06-03") == (
        "2026-06-01",
        "2026-06-30",
    )


def test_filter_records_keywords_and_export(tmp_path: Path, monkeypatch) -> None:
    import src.output.analytics as analytics
    import src.output.cumulative_exporter as cumulative_exporter

    monkeypatch.setattr(analytics, "DATA_ROOT", tmp_path / "data")
    monkeypatch.setattr(cumulative_exporter, "DATA_ROOT", tmp_path / "data")
    master_path = tmp_path / "master.docx"
    _save_master(master_path)
    _write_json(
        tmp_path / "data" / "output" / "cumulative_digest.json",
        {
            "date_blocks": [
                {
                    "date": "2026-06-02",
                    "items": [
                        {
                            "title_cn": "日本发布AI政策",
                            "url": "https://example.com/a",
                            "source": "Nikkei",
                            "type": "政策",
                            "soft_hard": "软科学",
                            "keywords": ["AI", "政策"],
                        }
                    ],
                }
            ]
        },
    )
    records = build_analytics_records(master_path)

    filtered = filter_records_by_date(records, "2026-06-01", "2026-06-03")
    keywords = keyword_counts(filtered, limit=5)
    output = export_analytics_excel(filtered, analytics_export_path("2026-06-01", "2026-06-03"))

    assert len(filtered) == 2
    assert "AI" in set(keywords["关键词"])
    assert output.exists()
