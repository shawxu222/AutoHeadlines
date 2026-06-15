from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.output.word_writer import digest_block_title, write_digest_docx


def test_write_digest_docx(tmp_path: Path) -> None:
    output = tmp_path / "digest.docx"
    write_digest_docx(
        [
            {
                "title_cn": "日本发布AI半导体计划",
                "summary_cn": "日本政府发布相关计划，推动研发基础设施建设。",
                "url": "https://example.com/news",
            }
        ],
        "2026-05-08",
        output,
    )
    assert output.exists()
    document = Document(output)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert digest_block_title("2026-05-08") in text
    assert "1.日本发布AI半导体计划" in text
    assert "" not in [paragraph.text for paragraph in document.paragraphs]
    assert document.styles["Normal"].font.size == Pt(12)
