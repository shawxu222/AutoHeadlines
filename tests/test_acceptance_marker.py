from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from src.output.acceptance_marker import (
    load_acceptance_entries,
    sync_acceptance_highlights,
)
from src.output.word_writer import add_digest_block


def _digests(*titles: str) -> list[dict[str, str]]:
    return [
        {
            "title_cn": title,
            "summary_cn": f"{title}摘要。",
            "url": f"https://example.com/{index}",
        }
        for index, title in enumerate(titles, start=1)
    ]


def _save_master(path: Path) -> None:
    document = Document()
    add_digest_block(document, "2026-06-02", _digests("已采纳新闻", "未采纳新闻"))
    add_digest_block(document, "2026-06-03", _digests("今日新闻"), leading_blank=True)
    document.paragraphs[1].runs[0].font.highlight_color = WD_COLOR_INDEX.YELLOW
    document.save(path)


def _append_manual_block(document: Document, heading: str, title: str, url: str) -> None:
    document.add_paragraph(heading)
    document.add_paragraph(f"1.{title}")
    document.add_paragraph(f"{title}摘要。")
    document.add_paragraph(url)


def _title_highlights(path: Path) -> dict[str, bool]:
    document = Document(path)
    output = {}
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if "." not in text:
            continue
        _, title = text.split(".", 1)
        if title.endswith("新闻"):
            output[title] = any(
                run.font.highlight_color == WD_COLOR_INDEX.YELLOW
                for run in paragraph.runs
            )
    return output


def test_load_acceptance_entries_reads_titles_and_existing_highlight(tmp_path: Path) -> None:
    master_path = tmp_path / "master.docx"
    _save_master(master_path)

    entries = load_acceptance_entries(master_path)

    assert [(entry.date, entry.order_index, entry.title) for entry in entries] == [
        ("2026-06-03", 1, "今日新闻"),
        ("2026-06-02", 1, "已采纳新闻"),
        ("2026-06-02", 2, "未采纳新闻"),
    ]
    assert [entry.accepted for entry in entries] == [False, True, False]
    assert entries[1].url == "https://example.com/1"


def test_load_acceptance_entries_handles_heading_variants(tmp_path: Path) -> None:
    master_path = tmp_path / "master.docx"
    document = Document()
    add_digest_block(document, "2025-07-22", _digests("正常格式新闻"))
    _append_manual_block(
        document,
        "每日科技要闻报送摘要20250723（旧组织）",
        "缺少冒号新闻",
        "https://example.com/723",
    )
    _append_manual_block(
        document,
        "每日科技要闻报送摘要20250827旧组织（编辑）：",
        "带署名新闻",
        "https://example.com/827",
    )
    document.save(master_path)

    entries = load_acceptance_entries(master_path)

    assert [(entry.date, entry.order_index, entry.title) for entry in entries] == [
        ("2025-08-27", 1, "带署名新闻"),
        ("2025-07-23", 1, "缺少冒号新闻"),
        ("2025-07-22", 1, "正常格式新闻"),
    ]


def test_sync_acceptance_highlights_updates_only_title_paragraphs(
    tmp_path: Path, monkeypatch
) -> None:
    import src.output.acceptance_marker as marker

    monkeypatch.setattr(marker, "DATA_ROOT", tmp_path / "data")
    master_path = tmp_path / "master.docx"
    _save_master(master_path)
    entries = load_acceptance_entries(master_path)
    accepted_ids = [
        entry.marker_id for entry in entries if entry.title in {"未采纳新闻", "今日新闻"}
    ]

    result = sync_acceptance_highlights(master_path, accepted_ids)

    assert result.item_count == 3
    assert result.accepted_count == 2
    assert result.backup_path is not None
    assert result.backup_path.exists()
    assert _title_highlights(master_path) == {
        "已采纳新闻": False,
        "未采纳新闻": True,
        "今日新闻": True,
    }
