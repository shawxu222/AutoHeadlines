from __future__ import annotations

from pathlib import Path

from docx import Document

from src.output.master_word_updater import update_master_digest_docx
from src.output.word_writer import add_digest_block, digest_block_title


def _sample_digest(title: str) -> list[dict[str, str]]:
    return [
        {
            "title_cn": title,
            "summary_cn": f"{title}的摘要。",
            "url": "https://example.com/news",
        }
    ]


def _docx_text(path: Path) -> str:
    return "\n".join(paragraph.text for paragraph in Document(path).paragraphs)


def test_update_master_digest_appends_new_date_block(tmp_path: Path) -> None:
    master_path = tmp_path / "master.docx"

    result = update_master_digest_docx(
        _sample_digest("日本发布AI政策"),
        "2026-06-03",
        master_path,
        backup=False,
    )

    assert result.action == "appended"
    assert result.backup_path is None
    text = _docx_text(master_path)
    assert digest_block_title("2026-06-03") in text
    assert "1.日本发布AI政策" in text


def test_update_master_digest_replaces_existing_date_block(tmp_path: Path) -> None:
    master_path = tmp_path / "master.docx"
    document = Document()
    add_digest_block(document, "2026-06-02", _sample_digest("旧新闻"))
    add_digest_block(
        document,
        "2026-06-03",
        _sample_digest("需要替换的旧标题"),
        leading_blank=True,
    )
    add_digest_block(document, "2026-06-04", _sample_digest("后一天新闻"), leading_blank=True)
    document.save(master_path)

    result = update_master_digest_docx(
        _sample_digest("替换后的新标题"),
        "2026-06-03",
        master_path,
        backup=False,
    )

    text = _docx_text(master_path)
    assert result.action == "replaced"
    assert text.count(digest_block_title("2026-06-03")) == 1
    assert "需要替换的旧标题" not in text
    assert "1.替换后的新标题" in text
    assert text.index(digest_block_title("2026-06-02")) < text.index(
        digest_block_title("2026-06-03")
    )
    assert text.index(digest_block_title("2026-06-03")) < text.index(
        digest_block_title("2026-06-04")
    )


def test_update_master_digest_finds_existing_legacy_heading_variants(
    tmp_path: Path,
) -> None:
    master_path = tmp_path / "master.docx"
    document = Document()
    document.add_paragraph("每日科技要闻报送摘要20260723（旧组织）")
    document.add_paragraph("1.需要替换的旧标题")
    document.add_paragraph("旧摘要。")
    document.add_paragraph("https://example.com/old")
    document.add_paragraph("每日科技要闻报送摘要20260724旧组织（编辑）：")
    document.add_paragraph("1.后一天新闻")
    document.add_paragraph("后一天摘要。")
    document.add_paragraph("https://example.com/next")
    document.save(master_path)

    result = update_master_digest_docx(
        _sample_digest("替换后的新标题"),
        "2026-07-23",
        master_path,
        backup=False,
    )

    text = _docx_text(master_path)
    assert result.action == "replaced"
    assert "需要替换的旧标题" not in text
    assert "1.替换后的新标题" in text
    assert "每日科技要闻报送摘要20260724旧组织（编辑）：" in text
    assert "1.后一天新闻" in text


def test_update_master_digest_creates_backup(tmp_path: Path, monkeypatch) -> None:
    import src.output.master_word_updater as updater

    monkeypatch.setattr(updater, "DATA_ROOT", tmp_path / "data")
    master_path = tmp_path / "master.docx"
    document = Document()
    add_digest_block(document, "2026-06-02", _sample_digest("已有新闻"))
    document.save(master_path)

    result = update_master_digest_docx(
        _sample_digest("新增新闻"),
        "2026-06-03",
        master_path,
        backup=True,
    )

    assert result.backup_path is not None
    assert result.backup_path.exists()
    assert "已有新闻" in _docx_text(result.backup_path)
