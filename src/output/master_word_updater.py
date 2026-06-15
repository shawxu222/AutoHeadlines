from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document

from src.config_loader import DATA_ROOT
from src.output.word_writer import add_digest_block
from src.utils.dates import compact_date


DATE_BLOCK_RE = re.compile(r"^每日科技要闻报送摘要\s*(?P<date>\d{8}).*$")


@dataclass(frozen=True)
class MasterWordUpdateResult:
    master_path: Path
    backup_path: Path | None
    action: str
    item_count: int


def update_master_digest_docx(
    digests: list[dict[str, Any]],
    run_date: str,
    master_path: Path,
    *,
    backup: bool = True,
    replace_existing: bool = True,
) -> MasterWordUpdateResult:
    """Append or replace a daily digest block inside the long-running master docx."""
    master_path = master_path.expanduser()
    if not str(master_path).strip():
        raise ValueError("总 Word 路径不能为空。")

    master_path.parent.mkdir(parents=True, exist_ok=True)
    backup_path = _backup_master_docx(master_path, run_date) if backup and master_path.exists() else None
    document = Document(master_path) if master_path.exists() else Document()

    existing_range = _find_digest_block(document, run_date)
    if existing_range and replace_existing:
        _replace_digest_block(document, existing_range, run_date, digests)
        action = "replaced"
    elif existing_range:
        add_digest_block(
            document,
            run_date,
            digests,
            leading_blank=_needs_leading_blank(document),
        )
        action = "appended_duplicate"
    else:
        add_digest_block(
            document,
            run_date,
            digests,
            leading_blank=_needs_leading_blank(document),
        )
        action = "appended"

    document.save(master_path)
    return MasterWordUpdateResult(
        master_path=master_path,
        backup_path=backup_path,
        action=action,
        item_count=len(digests),
    )


def _backup_master_docx(master_path: Path, run_date: str) -> Path:
    backup_dir = DATA_ROOT / "output" / "backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = backup_dir / f"master_digest_{_safe_date(run_date)}_{timestamp}.docx"
    shutil.copy2(master_path, backup_path)
    return backup_path


def _safe_date(run_date: str) -> str:
    return re.sub(r"\D", "", run_date) or datetime.now().strftime("%Y%m%d")


def _find_digest_block(document: Document, run_date: str) -> tuple[int, int] | None:
    paragraphs = document.paragraphs
    target_date = compact_date(run_date)
    start = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if _heading_compact_date(paragraph.text.strip()) == target_date
        ),
        None,
    )
    if start is None:
        return None

    end = len(paragraphs)
    for index in range(start + 1, len(paragraphs)):
        if _heading_compact_date(paragraphs[index].text.strip()):
            end = index
            break
    return start, end


def _heading_compact_date(text: str) -> str:
    date_match = DATE_BLOCK_RE.match(text)
    if not date_match:
        return ""
    return date_match.group("date")


def _replace_digest_block(
    document: Document,
    block_range: tuple[int, int],
    run_date: str,
    digests: list[dict[str, Any]],
) -> None:
    start, end = block_range
    paragraphs = document.paragraphs
    anchor = paragraphs[end]._element if end < len(paragraphs) else None
    for paragraph in paragraphs[start:end]:
        paragraph._element.getparent().remove(paragraph._element)

    before_count = len(document.paragraphs)
    add_digest_block(document, run_date, digests)
    new_elements = [paragraph._element for paragraph in document.paragraphs[before_count:]]
    if anchor is None:
        return

    for element in new_elements:
        element.getparent().remove(element)
        anchor.addprevious(element)


def _needs_leading_blank(document: Document) -> bool:
    return any(paragraph.text.strip() for paragraph in document.paragraphs)
