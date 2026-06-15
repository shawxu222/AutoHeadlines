from __future__ import annotations

import hashlib
import re
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from src.config_loader import DATA_ROOT
from src.parsers.text_cleaner import clean_text


DATE_BLOCK_RE = re.compile(r"^每日科技要闻报送摘要\s*(?P<date>\d{8}).*$")
ITEM_TITLE_RE = re.compile(r"^\s*(?P<index>\d+)[.．、]\s*(?P<title>.+?)\s*$")
URL_RE = re.compile(r"https?://\S+")


@dataclass(frozen=True)
class AcceptanceEntry:
    marker_id: str
    date: str
    order_index: int
    title: str
    url: str
    accepted: bool
    paragraph_index: int


@dataclass(frozen=True)
class AcceptanceSyncResult:
    master_path: Path
    backup_path: Path | None
    item_count: int
    accepted_count: int


def load_acceptance_entries(master_path: Path) -> list[AcceptanceEntry]:
    master_path = master_path.expanduser()
    if not master_path.exists():
        return []
    document = Document(master_path)
    return _collect_acceptance_entries(document)


def sync_acceptance_highlights(
    master_path: Path,
    accepted_marker_ids: Iterable[str],
    *,
    backup: bool = True,
) -> AcceptanceSyncResult:
    master_path = master_path.expanduser()
    if not master_path.exists():
        raise FileNotFoundError(f"找不到总 Word：{master_path}")

    accepted_set = {str(marker_id) for marker_id in accepted_marker_ids if marker_id}
    backup_path = _backup_master_docx(master_path) if backup else None
    document = Document(master_path)
    entries = _collect_acceptance_entries(document)
    paragraphs = document.paragraphs

    for entry in entries:
        selected = entry.marker_id in accepted_set
        _set_title_highlight(paragraphs[entry.paragraph_index], selected)

    document.save(master_path)
    return AcceptanceSyncResult(
        master_path=master_path,
        backup_path=backup_path,
        item_count=len(entries),
        accepted_count=sum(1 for entry in entries if entry.marker_id in accepted_set),
    )


def _collect_acceptance_entries(document: Document) -> list[AcceptanceEntry]:
    entries: list[AcceptanceEntry] = []
    current_date = ""
    paragraphs = document.paragraphs

    for index, paragraph in enumerate(paragraphs):
        text = _paragraph_text(paragraph)
        heading_date = _date_from_block_heading(text)
        if heading_date:
            current_date = heading_date
            continue
        if not current_date:
            continue

        item_match = ITEM_TITLE_RE.match(text)
        if not item_match:
            continue

        order_index = int(item_match.group("index"))
        title = clean_text(item_match.group("title"))
        if not title:
            continue
        url = _find_following_url(paragraphs, index + 1)
        accepted = _paragraph_has_yellow_highlight(paragraph)
        entries.append(
            AcceptanceEntry(
                marker_id=_marker_id(current_date, order_index, title, url),
                date=current_date,
                order_index=order_index,
                title=title,
                url=url,
                accepted=accepted,
                paragraph_index=index,
            )
        )
    return sorted(
        entries,
        key=lambda item: (item.date, -item.order_index, item.title),
        reverse=True,
    )


def _find_following_url(paragraphs, start: int) -> str:
    for paragraph in paragraphs[start : start + 4]:
        text = _paragraph_text(paragraph)
        if _date_from_block_heading(text) or ITEM_TITLE_RE.match(text):
            return ""
        match = URL_RE.search(text)
        if match:
            return match.group(0)
    return ""


def _marker_id(date_key: str, order_index: int, title: str, url: str) -> str:
    payload = f"{date_key}|{order_index}|{clean_text(title)}|{clean_text(url)}"
    return hashlib.sha1(payload.encode("utf-8")).hexdigest()[:16]


def _date_from_block_heading(text: str) -> str:
    date_match = DATE_BLOCK_RE.match(text)
    if not date_match:
        return ""
    return _iso_from_compact(date_match.group("date"))


def _set_title_highlight(paragraph, accepted: bool) -> None:
    for run in paragraph.runs:
        if accepted:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        elif _is_yellow_highlight(run.font.highlight_color):
            run.font.highlight_color = None


def _paragraph_has_yellow_highlight(paragraph) -> bool:
    return any(_is_yellow_highlight(run.font.highlight_color) for run in paragraph.runs)


def _is_yellow_highlight(value) -> bool:
    return (
        value == WD_COLOR_INDEX.YELLOW
        or getattr(value, "name", "") == "YELLOW"
        or "YELLOW" in str(value)
    )


def _paragraph_text(paragraph) -> str:
    try:
        text = "".join(node.text or "" for node in paragraph._element.xpath(".//w:t"))
        return clean_text(text)
    except Exception:
        return clean_text(paragraph.text)


def _iso_from_compact(value: str) -> str:
    try:
        return datetime.strptime(value, "%Y%m%d").date().isoformat()
    except ValueError:
        return value


def _backup_master_docx(master_path: Path) -> Path:
    backup_dir = DATA_ROOT / "output" / "backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = backup_dir / f"master_digest_acceptance_{timestamp}.docx"
    shutil.copy2(master_path, backup_path)
    return backup_path
