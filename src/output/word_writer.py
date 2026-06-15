from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from src.config_loader import DATA_ROOT, digest_title_template, load_profile
from src.utils.dates import compact_date


def final_docx_path(run_date: str) -> Path:
    return DATA_ROOT / "output" / f"final_digest_{compact_date(run_date)}.docx"


def cumulative_docx_path() -> Path:
    return DATA_ROOT / "output" / "cumulative_digest.docx"


def write_digest_docx(
    digests: list[dict[str, Any]], run_date: str, path: Path | None = None
) -> Path:
    output_path = path or final_docx_path(run_date)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _set_default_font(document)
    add_digest_block(document, run_date, digests)
    document.save(output_path)
    return output_path


def write_cumulative_docx(
    date_blocks: list[dict[str, Any]], path: Path | None = None
) -> Path:
    output_path = path or cumulative_docx_path()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _set_default_font(document)
    for index, block in enumerate(date_blocks):
        if index > 0:
            _apply_paragraph_format(document.add_paragraph())
            _apply_paragraph_format(document.add_paragraph())
        add_digest_block(document, str(block["date"]), block.get("items", []))

    document.save(output_path)
    return output_path


def digest_block_title(run_date: str) -> str:
    return digest_title_template().format(date=compact_date(run_date))


def add_digest_block(
    document: Document,
    run_date: str,
    digests: list[dict[str, Any]],
    *,
    leading_blank: bool = False,
) -> None:
    if leading_blank:
        _apply_paragraph_format(document.add_paragraph())

    title_paragraph = document.add_paragraph()
    _apply_paragraph_format(title_paragraph)
    title_run = title_paragraph.add_run(digest_block_title(run_date))
    _apply_run_font(title_run, bold=True)

    for index, digest in enumerate(digests, start=1):
        heading = document.add_paragraph()
        _apply_paragraph_format(heading)
        run = heading.add_run(f"{index}.{digest.get('title_cn', '')}")
        _apply_run_font(run, bold=True)

        summary = document.add_paragraph()
        _apply_paragraph_format(summary)
        summary_run = summary.add_run(str(digest.get("summary_cn", "")))
        _apply_run_font(summary_run)

        url = str(digest.get("url", "")).strip()
        url_paragraph = document.add_paragraph()
        _apply_paragraph_format(url_paragraph)
        if url:
            _add_hyperlink(url_paragraph, url, url)


def _set_default_font(document: Document) -> None:
    font = load_profile().word_font
    style = document.styles["Normal"]
    style.font.name = font
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    style.font.size = Pt(12)


def _apply_paragraph_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(20)
    paragraph.paragraph_format.space_after = Pt(0)


def _apply_run_font(run, bold: bool = False) -> None:
    font = load_profile().word_font
    run.bold = bold
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(12)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    font = load_profile().word_font
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    new_run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    run_properties.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(underline)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), font)
    fonts.set(qn("w:hAnsi"), font)
    fonts.set(qn("w:eastAsia"), font)
    run_properties.append(fonts)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "24")
    run_properties.append(size)

    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "24")
    run_properties.append(size_cs)

    new_run.append(run_properties)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
